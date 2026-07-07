"""
DOCX Generator — EXACT HAIRAT template replica.

Fonts, sizes, colors, margins, spacing — all extracted directly from the
HAIRAT Salesforce BA PDF using PyMuPDF font inspection.

Template style spec:
  Page:             A4, margins L=54pt R=46pt T=36pt B=36pt
  Contact row:      TimesNewRomanPSMT  11pt  black  (tab-separated L/R)
  Name:             TimesNewRomanPS-BoldMT  16pt  black  centered
  LinkedIn:         TimesNewRomanPSMT  11pt  black
  Specialisation:   TimesNewRomanPSMT  11pt  black  centered
  Section headers:  TimesNewRomanPS-BoldMT  12pt  #4472C4  + underline border
  Role/Client/Date: TimesNewRomanPS-BoldMT  11pt  black  bold
  Bullet ●:         Calibri  11pt  black
  Bullet text:      TimesNewRomanPSMT  11pt  black
"""

import io
import logging
from typing import Optional, List

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# ── Template colour from PDF inspection ───────────────────────────────────────
BLUE_R, BLUE_G, BLUE_B = 0x44, 0x72, 0xC4   # #4472C4 — section header colour

# ── Font names (exactly as they appear in the PDF) ────────────────────────────
FONT_TNR        = "Times New Roman"   # maps to TimesNewRomanPSMT
FONT_CALIBRI    = "Calibri"           # bullet ● character

# ── Sizes (from PDF span inspection) ─────────────────────────────────────────
SZ_CONTACT  = Pt(11)
SZ_NAME     = Pt(16)
SZ_SECTION  = Pt(12)
SZ_ROLE     = Pt(11)
SZ_BODY     = Pt(11)

# ── Margin (PDF L=54pt / R=46pt / T=36pt / B=36pt → convert to Inches) ───────
# 54 pt = 0.75 inch,  36 pt = 0.5 inch
MARGIN_LR = Inches(0.75)
MARGIN_TB = Inches(0.5)


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _set_run(run, font: str, size: Pt, bold=False, italic=False,
             color: Optional[tuple] = None, underline=False):
    run.font.name   = font
    run.font.size   = size
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if underline:
        run.font.underline = True


def _new_para(doc: Document,
              alignment=WD_ALIGN_PARAGRAPH.LEFT,
              space_before=0,
              space_after=0,
              line_spacing=None) -> "Paragraph":
    """Add a paragraph with explicit spacing."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment    = alignment
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing is not None:
        pf.line_spacing = Pt(line_spacing)
    return para


def _add_section_border(para):
    """
    Add a solid blue bottom border under the paragraph —
    this is how HAIRAT draws the line under every section header.
    """
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")            # 0.75pt line
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")       # exact template blue
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_tab_stop_right(para, position_inches: float):
    """Add a right-aligned tab stop at position_inches from left margin."""
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(position_inches * 1440)))  # twips
    tabs.append(tab)
    pPr.append(tabs)


# ── Section builders ──────────────────────────────────────────────────────────

def _add_contact_row(doc: Document, phone: str, email: str):
    """
    Phone:XXXXXXXXXX[right-tab]Email: email@domain
    Matches PDF span: timesNewRomanPSMT 11pt, tab-separated.
    """
    para = _new_para(doc, space_before=0, space_after=2)
    _set_tab_stop_right(para, 6.5)

    left_text  = f"Phone: {phone}" if phone else ""
    right_text = f"Email: {email}" if email else ""

    if left_text:
        r = para.add_run(left_text)
        _set_run(r, FONT_TNR, SZ_CONTACT)
    if left_text and right_text:
        r = para.add_run("\t")
        _set_run(r, FONT_TNR, SZ_CONTACT)
    if right_text:
        r = para.add_run(right_text)
        _set_run(r, FONT_TNR, SZ_CONTACT)


def _add_name(doc: Document, name: str):
    """
    Centred, TimesNewRomanPS-BoldMT 16pt bold — exactly as in PDF.
    """
    para = _new_para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=6, space_after=0)
    r = para.add_run(name)
    _set_run(r, FONT_TNR, SZ_NAME, bold=True)


def _add_linkedin_line(doc: Document, linkedin: str, specialization: str):
    """LinkedIn URL + (specialization) on the next line, centred."""
    if linkedin:
        para = _new_para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=2, space_after=0)
        prefix = para.add_run("LinkedIn: ")
        _set_run(prefix, FONT_TNR, SZ_BODY)
        url = para.add_run(linkedin)
        _set_run(url, FONT_TNR, SZ_BODY, underline=True,
                 color=(0x00, 0x70, 0xC0))  # hyperlink blue

    if specialization:
        para = _new_para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=0, space_after=4)
        r = para.add_run(f"({specialization})")
        _set_run(r, FONT_TNR, SZ_BODY)


def _add_section_header(doc: Document, title: str):
    """
    SECTION HEADER — TimesNewRomanPS-BoldMT 12pt #4472C4 bold + blue bottom border.
    Matches PDF exactly.
    """
    para = _new_para(doc, space_before=10, space_after=4)
    r = para.add_run(title.upper())
    _set_run(r, FONT_TNR, SZ_SECTION, bold=True,
             color=(BLUE_R, BLUE_G, BLUE_B))
    _add_section_border(para)


def _add_bullet(doc: Document, text: str, color=(0, 0, 0)):
    """
    ●[space]text  —  bullet in Calibri 11pt, text in TimesNewRoman 11pt.
    Left indent = 0.35" to match PDF's hanging-indent style.
    """
    para = _new_para(doc, space_before=0, space_after=2)
    pf = para.paragraph_format
    pf.left_indent   = Inches(0.25)
    pf.first_line_indent = Inches(-0.15)

    # Bullet character in Calibri (matches PDF)
    bullet_run = para.add_run("\u25cf ")
    _set_run(bullet_run, FONT_CALIBRI, SZ_BODY, color=color)

    # Body text in Times New Roman
    text_run = para.add_run(text.strip())
    _set_run(text_run, FONT_TNR, SZ_BODY, color=color)


def _add_role_line(doc: Document, text: str, space_before=8):
    """Bold Times New Roman 11pt — used for Role / Client / Date lines."""
    para = _new_para(doc, space_before=space_before, space_after=0)
    r = para.add_run(text)
    _set_run(r, FONT_TNR, SZ_ROLE, bold=True)


def _add_plain_line(doc: Document, text: str, space_before=2):
    """Regular Times New Roman 11pt — used for education, brief entries."""
    para = _new_para(doc, space_before=space_before, space_after=0)
    r = para.add_run(text)
    _set_run(r, FONT_TNR, SZ_BODY)


# ── Main entry point ──────────────────────────────────────────────────────────

def generate_docx(data: dict) -> bytes:
    """
    Build a DOCX byte string from the structured ConvertedResumeData dict.
    Style matches the HAIRAT Salesforce BA PDF template exactly:
      - Times New Roman 16pt bold centred name
      - Times New Roman 12pt bold #4472C4 section headers with bottom border
      - Calibri ● + Times New Roman 11pt body bullets
    """
    doc = Document()

    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Margins
    for section in doc.sections:
        section.top_margin    = MARGIN_TB
        section.bottom_margin = MARGIN_TB
        section.left_margin   = MARGIN_LR
        section.right_margin  = MARGIN_LR

    phone          = (data.get("phone")          or "").strip()
    email          = (data.get("email")          or "").strip()
    name           = (data.get("name")           or "Candidate").strip()
    linkedin       = (data.get("linkedin")       or "").strip()
    specialization = (data.get("specialization") or "").strip()

    # ── 1. Contact row ────────────────────────────────────────────────────────
    _add_contact_row(doc, phone, email)

    # ── 2. Name ───────────────────────────────────────────────────────────────
    _add_name(doc, name)

    # ── 3. LinkedIn + Specialization ──────────────────────────────────────────
    _add_linkedin_line(doc, linkedin, specialization)

    # ── 4. PROFESSIONAL SUMMARY ───────────────────────────────────────────────
    summary = data.get("professional_summary") or []
    if summary:
        _add_section_header(doc, "Professional Summary")
        for point in summary:
            _add_bullet(doc, point)

    # ── 5. TECHNICAL SKILLS ───────────────────────────────────────────────────
    skills = data.get("technical_skills") or []
    if skills:
        _add_section_header(doc, "Technical Skills")
        for skill in skills:
            _add_bullet(doc, skill)

    # ── 6. PROFESSIONAL EXPERIENCE ────────────────────────────────────────────
    experience = data.get("professional_experience") or []
    if experience:
        _add_section_header(doc, "Professional Experience")
        for idx, job in enumerate(experience):
            sb = 10 if idx == 0 else 8
            role   = job.get("role", "")
            client = job.get("client", "")
            dates  = job.get("date_range", "")

            _add_role_line(doc, f"Role: {role}", space_before=sb)
            _add_role_line(doc, f"Client: {client}", space_before=2)
            _add_role_line(doc, dates, space_before=2)

            for resp in (job.get("responsibilities") or []):
                _add_bullet(doc, resp)

    # ── 7. OTHER RELEVANT EXPERIENCE ─────────────────────────────────────────
    other = data.get("other_experience") or []
    if other:
        _add_section_header(doc, "Other Relevant Experience")
        for entry in other:
            role    = entry.get("role", "")
            company = entry.get("company", "")
            dates   = entry.get("date_range", "")
            _add_plain_line(doc, f"{role} at {company}", space_before=6)
            _add_plain_line(doc, dates, space_before=1)

    # ── 8. EDUCATION HISTORY ──────────────────────────────────────────────────
    education = data.get("education") or []
    if education:
        _add_section_header(doc, "Education History")
        for edu in education:
            _add_plain_line(doc, edu, space_before=6)

    # ── 9. CERTIFICATIONS ─────────────────────────────────────────────────────
    certs = data.get("certifications") or []
    if certs:
        _add_section_header(doc, "Certifications")
        for cert in certs:
            _add_bullet(doc, cert)

    buf = io.BytesIO()
    doc.save(buf)
    logger.info("DOCX generated — %d bytes", buf.tell())
    return buf.getvalue()
