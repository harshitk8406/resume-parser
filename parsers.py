"""
Resume file parsers for PDF, DOCX, and plain-text formats.
Returns raw extracted text for the AI agent to process.
"""

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF is required for PDF parsing. Run: pip install PyMuPDF")

    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page_num, page in enumerate(doc):
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(f"[Page {page_num + 1}]\n{page_text}")

    full_text = "\n\n".join(text_parts)
    if not full_text.strip():
        raise ValueError("No extractable text found in the PDF. The file may be image-based (scanned).")
    return full_text


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Run: pip install python-docx")

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    full_text = "\n".join(paragraphs)
    if not full_text.strip():
        raise ValueError("No text found in the DOCX file.")
    return full_text


def parse_txt(file_bytes: bytes) -> str:
    """Decode plain-text resume bytes."""
    for encoding in ("utf-8", "utf-16", "latin-1", "cp1252"):
        try:
            text = file_bytes.decode(encoding)
            if text.strip():
                return text.strip()
        except (UnicodeDecodeError, ValueError):
            continue
    raise ValueError("Unable to decode the text file with common encodings.")


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Dispatch to the correct parser based on file extension.

    Args:
        filename: Original filename (used to determine type).
        file_bytes: Raw file content.

    Returns:
        Extracted plain text from the resume.

    Raises:
        ValueError: If the file type is unsupported or text extraction fails.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        logger.info("Parsing PDF resume: %s", filename)
        return parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        logger.info("Parsing DOCX resume: %s", filename)
        return parse_docx(file_bytes)
    elif ext in ("txt", "text", ""):
        logger.info("Parsing TXT resume: %s", filename)
        return parse_txt(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type '.{ext}'. Please upload a PDF, DOCX, or TXT file."
        )
